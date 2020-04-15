import json

# DOCX or python-docx is used to create and save the Beautyreport as a Word document
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, Pt, RGBColor
from docx.text.tabstops import TabStop, TabStops

from html.parser import HTMLParser

def handle_word(form_data):
    # Get the raw form data which got generated out of the submitted anamnese form page
    # json.loads converts it to a dictionary
    content = json.loads(form_data)
    # from_data has two fields. A User ID gets sent to identify to whom the beautyreport data belongs and
    # a data field in which the entire generated text (plus user data) is stored as a JSON string.
    # raw_data saves the content of this data field in a seperate variable so that it can be processed further.
    raw_data = content['data']
    # As raw_data is a string in JSON format, it must also be converted to a dictionary
    data = json.loads(raw_data)

    # Open the template document
    document = Document("media/template.docx")
    # By not providing the path to an existing document, one could also initialize a new document.
    # Our beauty report is based on the existing document at media/template.docx, which is a MS Word document
    # which has already basic formatting applied, like the paper size of DIN-A5 and border margin.

    last_run_identifier = len(data) - 1

    for count, ka in enumerate(data):
        print("**** Word Debugging ****")
        print("*** Chapter:")
        print(data[ka])
        print("--------------------------------------------------")
        document.add_paragraph(data[ka]['chapterHeader'], style='L1 Heading')
        print("--------------------------------------------------")
        print("*** Adding Heading:")
        print(data[ka]['chapterHeader'])
        print("--------------------------------------------------")
        for kb in data[ka]:
            if not kb == 'chapterHeader':
                if not data[ka][kb]['subChapterHeader'] == '':
                    document.add_paragraph(data[ka][kb]['subChapterHeader'], style='L2 Heading')
                    print("--------------------------------------------------")
                    print("*** Adding Sub Heading:")
                    print(data[ka][kb]['subChapterHeader'])
                    print("--------------------------------------------------")
                for kc in data[ka][kb]:
                    if not kc == 'subChapterHeader':
                        print("--------------------------------------------------")
                        print("*** Adding Paragraph:")
                        print(data[ka][kb][kc])
                        print("--------------------------------------------------")
                        try:
                            parser = DocumentHTMLParser(document)
                            parser.add_paragraph_and_feed(data[ka][kb][kc]['text'])
                            print("* Sucessfully added Paragraph")
                            print("--------------------------------------------------")
                        except:
                            print("* Failed adding Paragraph")
                            print("--------------------------------------------------")
                            pass

        # Add a page break after each chapter expect the very last one
        lb = document.add_paragraph()
        if not count == last_run_identifier:
            lb.add_run().add_break(WD_BREAK.PAGE)

    return document


class DocumentHTMLParser(HTMLParser):
    def __init__(self, document):
        HTMLParser.__init__(self)
        self.document = document

    def add_paragraph_and_feed(self, html):
        self.paragraph = self.document.add_paragraph()
        self.run = self.paragraph.add_run()
        self.feed(html)

    def handle_starttag(self, tag, attrs):
        if tag == "b":
            # Text bold
            self.run.bold = True

        if tag == "i":
            # Text italic
            self.run.italic = True

        if tag == "u":
            # Text underlined
            self.run.underline = True

        if tag in ["br", "ul", "ol"]:
            pass

        if tag == "li":
            # Add a listing point at every list item
            self.run.add_text(u'● ')

        if tag == "tab":
            # Add a tab stop with <tab>
            self.run.add_tab()


    def handle_endtag(self, tag):
        if tag in ["br", "li", "ul", "ol"]:
            # Starting a new paragraph starts a new line, which is also wished within listings
            self.paragraph = self.document.add_paragraph()
        if tag == "p":
            self.paragraph = self.document.add_paragraph()

        # Add the paragraph to the current document's session
        self.run = self.paragraph.add_run()

    def handle_data(self, data):
        # Set font size
        self.run.font.size = Pt(12)

        # Set font family
        self.run.font.name = "Oxfam TSTAR PRO"

        # Set the paragraph alignment; This sets the alignment to justificaton
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Add the text to the current paragraph
        self.run.add_text(data)