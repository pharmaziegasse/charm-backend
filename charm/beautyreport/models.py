import json
import os
import pytz

from charm.coach.models import Coach

from django.contrib.auth import get_user_model
from django.core.serializers.json import DjangoJSONEncoder

from django.db import models

from django.conf import settings
from django.utils import timezone

# DOCX or python-docx is used to create and save the Beautyreport as a Word document
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, Pt, RGBColor
from docx.text.tabstops import TabStop, TabStops

from html.parser import HTMLParser

# GrabzIT is a library for converting HTML text to Word document formatting
# from GrabzIt import GrabzItClient
# from GrabzIt import GrabzItDOCXOptions

from modelcluster.fields import ParentalKey

from wagtail.admin.edit_handlers import FieldPanel, MultiFieldPanel, InlinePanel
from wagtail.contrib.forms.models import AbstractEmailForm, AbstractFormField, AbstractFormSubmission


class DocumentHTMLParser(HTMLParser):
    def __init__(self, document):
        HTMLParser.__init__(self)
        self.document = document
        # self.paragraph = self.document.add_paragraph()
        # self.run = self.paragraph.add_run()

    def add_paragraph_and_feed(self, html):
        self.paragraph = self.document.add_paragraph()
        self.run = self.paragraph.add_run()
        self.feed(html)

    def handle_starttag(self, tag, attrs):
        # self.run = self.paragraph.add_run()
        if tag == "b":
            # Text bold
            self.run.bold = True
        if tag == "i":
            # Text italic
            self.run.italic = True
        if tag == "u":
            # Text underlined
            self.run.underline = True
        if tag in ["br", "ul", "ol"]:
            # self.run.add_break()
            pass
        if tag == "li":
            # Add a listing point at every list item
            self.run.add_text(u'● ')
        if tag == "tab":
            # Adds a tab stop with <tab>
            self.run.add_tab()
        # if tag == "table":
        #     # Adds a table with <table>
        #     table = self.document.add_table()
        #     table_active = True
        #     # row = table.rows[0]
        #     # row.cells[0].text = "test"
        # if tag == "row":
        #     if table_active:
        #         table.add_row()
        #     # Adds a row to a table
        # if tag == "col":
        #     if table_active:
        #         table.add_column()
        # global table
        # table = None
        # global table_active
        # table_active = False
        # global table_row 
        # table_row = 0
        # global table_cell 
        # table_cell = 0



    def handle_endtag(self, tag):
        if tag in ["br", "li", "ul", "ol"]:
            # self.run.add_break()
            self.paragraph = self.document.add_paragraph()
        if tag == "p":
            self.paragraph = self.document.add_paragraph()
        # if tag == "table":
        #     table_active = False
        self.run = self.paragraph.add_run()

    def handle_data(self, data):
        self.run.font.size = Pt(12)
        self.run.font.name = "Oxfam TSTAR PRO"
        self.paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.run.add_text(data)
        
    # def handle_starttag(self, tag, attrs):
    #     print("Encountered a start tag:", tag)

    # def handle_endtag(self, tag):
    #     print("Encountered an end tag :", tag)

    # def handle_data(self, data):
    #     print("Encountered some data  :", data)

class BeautyreportDocument(models.Model):
    # beautyreport = models.ForeignKey(
    #     "Beautyreport", null=True,
    #     on_delete=models.SET_NULL
    # )
    link = models.CharField(
        null=True, blank=True,
        max_length=256
    )

class FormField(AbstractFormField):
    page = ParentalKey('BrFormPage', on_delete=models.CASCADE, related_name='form_fields')

class Beautyreport(models.Model):
    date = models.DateTimeField(
        null=True, blank=True
    )
    # This field identifies to which user the anamnese report belongs to
    user = models.ForeignKey(
        get_user_model(), null=True,
        on_delete=models.SET_NULL
    )
    # To identify which coach created this beautyreport
    coach = models.ForeignKey(
        Coach, null=True,
        on_delete=models.SET_NULL,
        related_name='Coach_BR'
    )
    document = models.ForeignKey(
        BeautyreportDocument,
        null=True, blank=True,
        on_delete=models.SET_NULL,
        related_name='Document_BR'
    )
    # docx_url = models.CharField(
    #     null=True, blank=True,
    #     max_length=128
    # )
    form_data = models.TextField(
        null=True, blank=True
    )

    # custom save function
    def save(self, *args, **kwargs):
        today = timezone.now()

        if not self.date:
            self.date = today


        '''
        Here starts the Beautyreport generation part.
        As this is a pretty complex part, it should be split up later.

        Current version: 1.1.0, 10/2/2019
        '''

        # Get the raw form data which got generated out of the submitted anamnese form page
        # json.loads converts it to a dictionary
        content = json.loads(self.form_data)
        # from_data has two fields. A User ID gets sent to identify to whom the beautyreport data belongs and
        # a data field in which the entire generated text (plus user data) is stored as a JSON string.
        # raw_data saves the content of this data field in a seperate variable so that it can be processed further.
        raw_data = content['data']
        # As raw_data is a JSON string again, it must also be converted to a dictionary
        data = json.loads(raw_data)

        # Initialize a new document
        document = Document("media/template.docx")
        # document = Document()

        '''
        # Get default document styles
        styles = document.styles

        # Style for Level 1 Heading
        heading1_style = styles.add_style('L1 Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.base_style = styles['List Number']
        font = heading1_style.font
        font.size = Pt(16)
        font.bold = True
        font.color.rgb = RGBColor(52, 90, 138)
        font.name = "Oxfam TSTAR PRO Bold"

        # Style for Level 2 Heading
        heading2_style = styles.add_style('L2 Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.base_style = styles['Heading 2']
        font = heading2_style.font
        font.size = Pt(13)
        font.name = "Oxfam TSTAR PRO Bold"

        # Paragraph Style (spacing)
        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.space_after = 0
        '''

        # document.save("template.docx")

        last_run_identifier = len(data) - 1

        for count, ka in enumerate(data):
            print("**** Word Debugging ****")
            print("*** Chapter:")
            print(data[ka])
            print("--------------------------------------------------")
            #document.add_heading(str(idx) + " " + data[ka]['chapterHeader'])
            document.add_paragraph(data[ka]['chapterHeader'], style='L1 Heading')
            print("--------------------------------------------------")
            print("*** Adding Heading:")
            print(data[ka]['chapterHeader'])
            print("--------------------------------------------------")
            for kb in data[ka]:
                if not kb == 'chapterHeader':
                    if not data[ka][kb]['subChapterHeader'] == '':
                        document.add_paragraph(data[ka][kb]['subChapterHeader'], style='L2 Heading')
                        print("--------------------------------------------------")
                        print("*** Adding Sub Heading:")
                        print(data[ka][kb]['subChapterHeader'])
                        print("--------------------------------------------------")
                    for kc in data[ka][kb]:
                        if not kc == 'subChapterHeader':
                            print("--------------------------------------------------")
                            print("*** Adding Paragraph:")
                            print(data[ka][kb][kc])
                            print("--------------------------------------------------")
                            try:
                                parser = DocumentHTMLParser(document)
                                parser.add_paragraph_and_feed(data[ka][kb][kc]['text'])
                                print("* Sucessfully added Paragraph")
                                print("--------------------------------------------------")
                            except:
                                print("* Failed adding Paragraph")
                                print("--------------------------------------------------")
                                pass
            # Page break after main chapter
            lb = document.add_paragraph()
            if not count == last_run_identifier:
                lb.add_run().add_break(WD_BREAK.PAGE)

        try:
            nid = BeautyreportDocument.objects.latest('id').id
        except:
            nid = 0

        document_name = 'Beautyreport_' \
            + self.user.first_name + '-' \
            + self.user.last_name + '_' \
            + today.strftime("%d-%m-%Y") \
            + '.docx'

        directory = settings.BR_DOCUMENT_PATH
        if not os.path.exists(directory):
            os.makedirs(directory)

        path = directory + document_name

        document.save(path)

        '''
        End of Beautyreport generation.
        '''

        blinkcollection = BeautyreportDocument(
            link=path
        )

        blinkcollection.save()

        self.document = blinkcollection

        super(Beautyreport, self).save(*args, **kwargs)

    class Meta:
        get_latest_by = "date"

class BrFormPage(AbstractEmailForm):
    content_panels = AbstractEmailForm.content_panels + [
        MultiFieldPanel(
            [
                InlinePanel('form_fields', label="Beautyreport fields")
            ],
            heading="Form Fields",
        )
    ]

    def get_submission_class(self):
        return BeautyreportFormSubmission

    def create_br(self, date, user, coach, form_data):
        br = Beautyreport(
            date = date,
            user = user,
            coach = coach,
            form_data = form_data
        )

        br.save()

        return br

    def process_form_submission(self, form):
        br = self.create_br(
            date = timezone.now(),
            user = get_user_model().objects.get(id=form.cleaned_data['uid']),
            coach = form.user,
            form_data = json.dumps(form.cleaned_data, cls=DjangoJSONEncoder)
        )

        self.get_submission_class().objects.create(
            form_data = json.dumps(form.cleaned_data, cls=DjangoJSONEncoder),
            page = self,
            br = br,
        )

class BeautyreportFormSubmission(AbstractFormSubmission):
    br = models.ForeignKey(Beautyreport, on_delete=models.CASCADE)
